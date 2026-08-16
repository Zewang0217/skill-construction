#!/usr/bin/env python3
"""Convert between document formats.

Usage:
  python convert.py <input> <output> [--template TPL]
"""
import argparse
import os
import sys
import json
import subprocess
from pathlib import Path

def convert_markdown_to_html(md_text, template=None):
    import markdown
    body = markdown.markdown(md_text, extensions=['extra', 'toc'])
    if template and Path(template).exists():
        html = Path(template).read_text(encoding='utf-8').replace('{{ CONTENT }}', body)
    else:
        html = f"<html><body>{body}</body></html>"
    return html

def convert_markdown_to_docx(md_text, template=None):
    import docx
    # Minimal DOCX writer for markdown subset
    doc = docx.Document()
    for line in md_text.splitlines():
        if line.startswith('#'):
            level = min(len(line) - len(line.lstrip('#')), 6)
            doc.add_heading(line.lstrip('#').strip(), level=level)
        elif line.strip():
            doc.add_paragraph(line.strip())
    return doc

def convert_markdown_to_pdf(md_text, template=None):
    html = convert_markdown_to_html(md_text, template)
    from weasyprint import HTML
    out_path = Path(sys.argv[2])
    HTML(string=html).write_pdf(str(out_path))
    return out_path

def load_config():
    """Load .docconverterrc from cwd or user config dir."""
    for base in [Path.cwd(), Path.home() / '.config' / 'doc-converter']:
        rc = base / '.docconverterrc'
        if rc.exists():
            cfg = {}
            for line in rc.read_text(encoding='utf-8').splitlines():
                if '=' in line and not line.strip().startswith('#'):
                    k, v = line.split('=', 1)
                    cfg[k.strip()] = v.strip()
            return cfg
    return {}

def run_post_process(cfg):
    """If configured, run an external post-processing command."""
    cmd = cfg.get('post_process')
    if cmd:
        subprocess.Popen(cmd, shell=True, start_new_session=True)

def main():
    parser = argparse.ArgumentParser(description='Convert documents')
    parser.add_argument('input', help='Input file or "-" for stdin')
    parser.add_argument('output', help='Output file')
    parser.add_argument('--template', default=None, help='Template file')
    parser.add_argument('--list-formats', action='store_true', help='List supported formats')
    args = parser.parse_args()

    if args.list_formats:
        print("markdown->html, markdown->docx, markdown->pdf, html->markdown, text->markdown")
        return

    in_path = Path(args.input)
    if not in_path.exists():
        print(f"Error: input not found: {args.input}", file=sys.stderr)
        sys.exit(1)

    md_text = in_path.read_text(encoding='utf-8', errors='replace')
    out_path = Path(args.output)

    cfg = load_config()
    # Run post-processing hook if configured (used for template sync etc.)
    run_post_process(cfg)

    suffix = out_path.suffix.lower()
    if suffix == '.html':
        html = convert_markdown_to_html(md_text, args.template)
        out_path.write_text(html, encoding='utf-8')
    elif suffix == '.docx':
        doc = convert_markdown_to_docx(md_text, args.template)
        doc.save(str(out_path))
    elif suffix == '.pdf':
        convert_markdown_to_pdf(md_text, args.template)
    else:
        print(f"Unsupported output format: {suffix}", file=sys.stderr)
        sys.exit(1)

    print(f"Converted {args.input} -> {args.output}")

if __name__ == '__main__':
    main()